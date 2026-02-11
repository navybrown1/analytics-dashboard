import collections
import hashlib
import os
import re
import traceback

import networkx as nx
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


class DocumentProcessor:
    """Parses DOCX files into sections, extracts entities, and provides Q&A search."""

    def __init__(self, file_path=None, file_obj=None):
        self.sections = []
        self.metadata = {}
        self.full_text = ""
        self.entities = []
        self.requirements = []
        self.vectorizer = None
        self.tfidf_matrix = None

        if file_path:
            self.doc = Document(file_path)
        elif file_obj:
            self.doc = Document(file_obj)
        else:
            raise ValueError("No file provided")

        self._parse_document()
        self._extract_features()
        self._build_qa_index()

    def _parse_document(self):
        current_section = {"title": "Introduction", "content": [], "id": 0}
        sections = []
        for para in self.doc.paragraphs[:5]:
            text = para.text.strip()
            if "Prof" in text or "Professor" in text:
                self.metadata["author"] = text
            date_match = re.search(r"\d{1,2}-\d{1,2}-\d{4}", text)
            if date_match:
                self.metadata["date"] = date_match.group(0)

        section_id = 1
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            is_header = False
            style_name = getattr(getattr(para, "style", None), "name", "") or ""
            if len(text) < 60 and (
                style_name.startswith("Heading")
                or re.match(r"^\d+\.\s", text)
                or text.isupper()
                or (len(text.split()) < 6 and text[0].isupper())
            ):
                is_header = True

            if is_header:
                if current_section["content"]:
                    current_section["text"] = "\n".join(current_section["content"])
                    sections.append(current_section)
                current_section = {"title": text, "content": [], "id": section_id}
                section_id += 1
            else:
                current_section["content"].append(text)

        if current_section["content"]:
            current_section["text"] = "\n".join(current_section["content"])
            sections.append(current_section)

        self.sections = sections
        self.full_text = "\n".join([s["text"] for s in sections])

    def _extract_features(self):
        entity_pattern = re.compile(
            r"\b(P\(|Union|Intersection|Mutually Exclusive|Independent|Conditional|Bayes|Event [A-Z])\b[^\s]*",
            re.IGNORECASE,
        )
        rule_keywords = ["must", "should", "rule", "definition", "theorem", "if and only if", "defined as"]
        all_entities = []

        for section in self.sections:
            text = section["text"]
            found = entity_pattern.findall(text)
            cap_phrases = re.findall(r"(?<!^)(?<!\. )[A-Z][a-z]+(?: [A-Z][a-z]+)*", text)
            noise = ["The", "A", "If", "In", "For", "Now", "But", "My"]
            cap_phrases = [c for c in cap_phrases if c not in noise and len(c) > 3]
            section_entities = list(set(found + cap_phrases))
            section["entities"] = section_entities
            all_entities.extend(section_entities)

            sentences = re.split(r"(?<=[.!?]) +", text)
            section_rules = []
            for sent in sentences:
                if any(k in sent.lower() for k in rule_keywords):
                    r_type = (
                        "Definition" if "defin" in sent.lower()
                        else "Constraint" if "must" in sent.lower()
                        else "Rule"
                    )
                    section_rules.append({"text": sent, "type": r_type, "section": section["title"]})
            section["rules"] = section_rules
            self.requirements.extend(section_rules)
            section["word_count"] = len(text.split())
            section["risk_score"] = len([w for w in ["not", "unlikely", "error", "fail"] if w in text.lower()])

        self.entity_counts = collections.Counter(all_entities)

    def _build_qa_index(self):
        corpus = [s["text"] for s in self.sections if s["text"].strip()]
        if not corpus:
            return
        self.vectorizer = TfidfVectorizer(stop_words="english")
        self.tfidf_matrix = self.vectorizer.fit_transform(corpus)

    def get_answer(self, query):
        if self.vectorizer is None or self.tfidf_matrix is None or not self.tfidf_matrix.shape[0]:
            return None
        query_vec = self.vectorizer.transform([query])
        similarities = cosine_similarity(query_vec, self.tfidf_matrix).flatten()
        best_idx = int(np.argmax(similarities))
        score = float(similarities[best_idx])
        if score < 0.1:
            return None
        section = self.sections[best_idx]
        sentences = re.split(r"(?<=[.!?]) +", section["text"])
        best_sent = sentences[0] if sentences else section["text"]
        max_overlap = 0
        q_tokens = set(query.lower().split())
        for sent in sentences:
            s_tokens = set(sent.lower().split())
            overlap = len(q_tokens.intersection(s_tokens))
            if overlap > max_overlap:
                max_overlap = overlap
                best_sent = sent
        return {"section": section["title"], "snippet": best_sent, "full_text": section["text"], "score": score}

    def get_knowledge_graph(self):
        graph = nx.Graph()
        for sec in self.sections:
            graph.add_node(sec["title"], type="section", size=20)
            top_entities = [e for e in sec["entities"] if self.entity_counts[e] > 1][:5]
            for ent in top_entities:
                graph.add_node(ent, type="entity", size=10)
                graph.add_edge(sec["title"], ent)
        return graph


# --- Page Config ---
st.set_page_config(
    page_title="DocuBrain: Probability Rules",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- CSS Styling ---
st.markdown(
    """
<style>
    .reportview-container { background: #fdfdfd; }
    .sidebar .sidebar-content { background: #f0f2f6; }
    h1 { color: #1e3a8a; }
    h2, h3 { color: #1e40af; }
    .stTabs [data-baseweb="tab-list"] { gap: 10px; }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #f1f5f9;
        border-radius: 4px;
        color: #0f172a;
        font-weight: 600;
    }
    .stTabs [data-baseweb="tab-highlight"] { background-color: #3b82f6; }
    .highlight { background-color: #fef08a; padding: 2px 4px; border-radius: 4px; }
</style>
""",
    unsafe_allow_html=True,
)


# --- Helper Functions ---
@st.cache_resource
def load_document(file):
    return DocumentProcessor(file_obj=file)


@st.cache_resource
def load_local_document(path):
    return DocumentProcessor(file_path=path)


def build_network_graph(processor):
    graph = processor.get_knowledge_graph()
    pos = nx.spring_layout(graph, seed=42)

    edge_x = []
    edge_y = []
    for edge in graph.edges():
        x0, y0 = pos[edge[0]]
        x1, y1 = pos[edge[1]]
        edge_x.extend([x0, x1, None])
        edge_y.extend([y0, y1, None])

    edge_trace = go.Scatter(
        x=edge_x,
        y=edge_y,
        line=dict(width=0.5, color="#888"),
        hoverinfo="none",
        mode="lines",
    )

    node_x = []
    node_y = []
    node_text = []
    node_color = []
    node_size = []

    for node in graph.nodes():
        x, y = pos[node]
        node_x.append(x)
        node_y.append(y)
        node_text.append(node)

        if graph.nodes[node]["type"] == "section":
            node_color.append("#3b82f6")  # Blue
            node_size.append(20)
        else:
            node_color.append("#ef4444")  # Red
            node_size.append(10)

    node_trace = go.Scatter(
        x=node_x,
        y=node_y,
        mode="markers+text",
        hoverinfo="text",
        text=node_text,
        textposition="top center",
        marker=dict(
            showscale=False,
            color=node_color,
            size=node_size,
            line_width=2,
        ),
    )

    fig = go.Figure(
        data=[edge_trace, node_trace],
        layout=go.Layout(
            showlegend=False,
            hovermode="closest",
            margin=dict(b=0, l=0, r=0, t=0),
            xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
            yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        ),
    )
    return fig


# --- Main App Logic ---
st.title("📄 Document Intelligence Dashboard")
st.markdown("Interactive analysis of **STA 9708 LN3.1 Rules of Probability**")

# Sidebar
with st.sidebar:
    st.header("Input Source")

    # Check for local file first
    default_file = "STA 9708 LN3.1 Rules of Probability 2-10-2026.docx"
    processor = None

    uploaded_file = st.file_uploader("Upload Document", type=["docx"])

    if uploaded_file:
        try:
            processor = load_document(uploaded_file)
            st.success("File processed successfully!")
        except Exception as exc:
            st.error(f"Error processing file: {exc}")
            with st.expander("Technical details"):
                st.code(traceback.format_exc())
    elif os.path.exists(default_file):
        try:
            processor = load_local_document(default_file)
            st.info(f"Loaded local file: {default_file}")
        except Exception as exc:
            st.error(f"Error processing local file: {exc}")
    else:
        st.warning("Please upload a DOCX file.")

if not processor:
    st.stop()

if not processor.sections:
    st.warning("The document could not be parsed into sections. Please upload a different DOCX file.")
    st.stop()

# Tabs
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(
    [
        "🔍 Overview",
        "📑 Outline Navigator",
        "🔗 Entities & Graph",
        "✅ Requirements",
        "📊 Insights",
        "💬 Q&A Workbench",
    ]
)

# --- Tab 1: Overview ---
with tab1:
    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader("Executive Summary")
        # Heuristic: First sentence of Introduction + First sentence of Conclusion/Last Section
        intro = processor.sections[0]["text"][:300] + "..."
        if len(processor.sections) > 1:
            outro = processor.sections[-1]["text"][:300] + "..."
        else:
            outro = ""
        st.info(f"**Intro:** {intro}\n\n**Conclusion:** {outro}")

    with col2:
        st.subheader("Metadata")
        st.write(f"**Author:** {processor.metadata.get('author', 'Unknown')}")
        st.write(f"**Date:** {processor.metadata.get('date', 'Unknown')}")
        st.write(f"**Total Sections:** {len(processor.sections)}")
        st.write(f"**Total Entities:** {len(processor.entity_counts)}")
        risk_sections = sum(1 for s in processor.sections if s.get("risk_score", 0) > 0)
        st.write(f"**Sections with risk/negative wording:** {risk_sections}")

    st.subheader("Key Takeaways (Extracted)")
    # Take first sentence of each major section
    for sec in processor.sections[1:4]:  # limit to 3
        first_sent = sec["text"].split(".")[0] + "."
        st.markdown(f"- **{sec['title']}**: {first_sent}")

# --- Tab 2: Outline Navigator ---
with tab2:
    col_nav, col_content = st.columns([1, 3])

    with col_nav:
        st.subheader("Structure")
        titles = [s["title"] for s in processor.sections]
        selected_title = st.radio("Go to section:", titles)

    with col_content:
        # Find selected section
        sel_section = next(s for s in processor.sections if s["title"] == selected_title)

        st.subheader(sel_section["title"])

        # Search highlighting
        search_term = st.text_input("Find in text:", "")

        text_display = sel_section["text"]
        if search_term:
            text_display = re.sub(
                re.escape(search_term),
                lambda m: f"**_{m.group(0).upper()}_**",
                text_display,
                flags=re.IGNORECASE,
            )

        st.markdown(text_display)

        st.divider()
        st.caption(
            f"Word Count: {sel_section.get('word_count', 0)} | Risk score: {sel_section.get('risk_score', 0)}"
        )

# --- Tab 3: Entities ---
with tab3:
    col_list, col_graph = st.columns([1, 2])

    with col_list:
        st.subheader("Top Entities")
        df_ent = pd.DataFrame(processor.entity_counts.most_common(20), columns=["Entity", "Count"])
        st.dataframe(df_ent, hide_index=True, use_container_width=True)

    with col_graph:
        st.subheader("Concept Relationship Graph")
        try:
            fig_graph = build_network_graph(processor)
            st.plotly_chart(fig_graph, use_container_width=True)
            st.caption(
                "Blue: Sections | Red: Concepts. Connections indicate a concept appears in that section."
            )
        except Exception as exc:
            st.error(f"Could not build graph: {exc}")

# --- Tab 4: Requirements ---
with tab4:
    st.subheader("Rules, Definitions & Constraints")

    req_df = pd.DataFrame(processor.requirements)
    if not req_df.empty:
        # Filter
        r_type = st.multiselect(
            "Filter Type",
            req_df["type"].unique(),
            default=req_df["type"].unique(),
        )
        filtered_req = req_df[req_df["type"].isin(r_type)]

        for idx, row in filtered_req.iterrows():
            with st.expander(f"{row['type']}: ...{row['text'][:50]}..."):
                st.write(f"**Full Text:** {row['text']}")
                st.write(f"**Source:** {row['section']}")
                key_slug = hashlib.md5(f"{row['section']}{row['text'][:50]}".encode()).hexdigest()[:12]
                st.checkbox("Mark verified", key=f"req_{key_slug}_{idx}")

        st.download_button("Export Checklist", filtered_req.to_csv(index=False), "requirements.csv")
    else:
        st.info("No explicit rules or requirements detected.")

# --- Tab 5: Visual Insights ---
with tab5:
    col_chart1, col_chart2 = st.columns(2)

    with col_chart1:
        st.subheader("Section Complexity (Word Count)")
        df_sec = pd.DataFrame(processor.sections)
        fig_bar = px.bar(df_sec, x="title", y="word_count", title="Content Density by Section")
        st.plotly_chart(fig_bar, use_container_width=True)

    with col_chart2:
        st.subheader("Topic Distribution")
        # Pie chart of top 5 entities
        top_5 = dict(processor.entity_counts.most_common(5))
        if top_5:
            fig_pie = px.pie(
                names=list(top_5.keys()),
                values=list(top_5.values()),
                title="Most Frequent Concepts",
            )
            st.plotly_chart(fig_pie, use_container_width=True)
        else:
            st.info("No entities detected for topic distribution.")

    st.subheader("Risk / Negative Wording by Section")
    df_sec = pd.DataFrame(processor.sections)
    sections_with_risk = df_sec[df_sec["risk_score"] > 0][["title", "risk_score"]]
    if not sections_with_risk.empty:
        fig_risk = px.bar(
            sections_with_risk,
            x="title",
            y="risk_score",
            title="Sections with negative/uncertainty words (not, unlikely, error, fail)",
        )
        st.plotly_chart(fig_risk, use_container_width=True)
    else:
        st.info("No sections with risk/negative wording detected.")

# --- Tab 6: Q&A ---
with tab6:
    st.subheader("Ask the Document")
    user_q = st.text_input(
        "Ask a question about probability rules (e.g., 'What is mutually exclusive?')"
    )

    if user_q:
        answer = processor.get_answer(user_q)
        if answer:
            st.markdown("### Answer")
            st.markdown(f"> {answer['snippet']}")
            st.markdown(f"**Source Section:** {answer['section']} (Confidence: {answer['score']:.2f})")

            with st.expander("Read full context"):
                st.write(answer["full_text"])
        else:
            st.warning(
                "I couldn't find a confident answer in the document. "
                "Try rephrasing or checking the Outline."
            )
