import collections
import re

import networkx as nx
import numpy as np
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


class DocumentProcessor:
    def __init__(self, file_path=None, file_obj=None):
        self.sections = []
        self.metadata = {}
        self.full_text = ""
        self.entities = []
        self.requirements = []
        self.vectorizer = None
        self.tfidf_matrix = None

        if file_path:
            self.doc = Document(file_path)
        elif file_obj:
            self.doc = Document(file_obj)
        else:
            raise ValueError("No file provided")

        self._parse_document()
        self._extract_features()
        self._build_qa_index()

    def _parse_document(self):
        """
        Parses the DOCX file into logical sections based on formatting heuristics.
        """
        current_section = {"title": "Introduction", "content": [], "id": 0}
        sections = []

        # Heuristics for metadata extraction from first few lines
        for para in self.doc.paragraphs[:5]:
            text = para.text.strip()
            if "Prof" in text or "Professor" in text:
                self.metadata["author"] = text
            date_match = re.search(r"\d{1,2}-\d{1,2}-\d{4}", text)
            if date_match:
                self.metadata["date"] = date_match.group(0)

        section_id = 1
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Heuristic: Short, bold lines or numbered lines are headers
            is_header = False
            if len(text) < 60 and (
                para.style.name.startswith("Heading")
                or re.match(r"^\d+\.\s", text)
                or text.isupper()
                or (len(text.split()) < 6 and text[0].isupper())
            ):
                is_header = True

            if is_header:
                # Save previous section
                if current_section["content"]:
                    current_section["text"] = "\n".join(current_section["content"])
                    sections.append(current_section)

                # Start new section
                current_section = {
                    "title": text,
                    "content": [],
                    "id": section_id,
                }
                section_id += 1
            else:
                current_section["content"].append(text)

        # Append final section
        if current_section["content"]:
            current_section["text"] = "\n".join(current_section["content"])
            sections.append(current_section)

        self.sections = sections
        self.full_text = "\n".join([s["text"] for s in sections])

    def _extract_features(self):
        """
        Extracts entities, requirements, and computes stats.
        """
        # 1. Regex definitions
        # Math/Prob terms (P(A), Union, etc)
        entity_pattern = re.compile(
            r"\b(P\(|Union|Intersection|Mutually Exclusive|Independent|Conditional|Bayes|Event [A-Z])\b[^\s]*",
            re.IGNORECASE,
        )
        # Requirements/Rules
        rule_keywords = [
            "must",
            "should",
            "rule",
            "definition",
            "theorem",
            "if and only if",
            "defined as",
        ]

        all_entities = []

        for section in self.sections:
            text = section["text"]

            # --- Entity Extraction ---
            found = entity_pattern.findall(text)
            # Add capitalized phrases that aren't start of sentences (simple NER)
            cap_phrases = re.findall(
                r"(?<!^)(?<!\. )[A-Z][a-z]+(?: [A-Z][a-z]+)*", text
            )

            # Filter common noise
            noise = ["The", "A", "If", "In", "For", "Now", "But", "My"]
            cap_phrases = [c for c in cap_phrases if c not in noise and len(c) > 3]

            section_entities = list(set(found + cap_phrases))
            section["entities"] = section_entities
            all_entities.extend(section_entities)

            # --- Requirements / Rules ---
            sentences = re.split(r"(?<=[.!?]) +", text)
            section_rules = []
            for sent in sentences:
                if any(k in sent.lower() for k in rule_keywords):
                    # Classify
                    r_type = (
                        "Definition"
                        if "defin" in sent.lower()
                        else "Constraint"
                        if "must" in sent.lower()
                        else "Rule"
                    )
                    section_rules.append(
                        {
                            "text": sent,
                            "type": r_type,
                            "section": section["title"],
                        }
                    )

            section["rules"] = section_rules
            self.requirements.extend(section_rules)

            # --- Stats ---
            section["word_count"] = len(text.split())
            section["risk_score"] = len(
                [w for w in ["not", "unlikely", "error", "fail"] if w in text.lower()]
            )

        # Frequency Count
        self.entity_counts = collections.Counter(all_entities)

    def _build_qa_index(self):
        """
        Builds a TF-IDF index for the Q&A system.
        """
        corpus = [s["text"] for s in self.sections if s["text"].strip()]
        if not corpus:
            return

        self.vectorizer = TfidfVectorizer(stop_words="english")
        self.tfidf_matrix = self.vectorizer.fit_transform(corpus)

    def get_answer(self, query):
        """
        Returns the most relevant section and a snippet for a query.
        """
        if self.vectorizer is None or self.tfidf_matrix is None or not self.tfidf_matrix.shape[0]:
            return None

        query_vec = self.vectorizer.transform([query])
        similarities = cosine_similarity(query_vec, self.tfidf_matrix).flatten()

        best_idx = np.argmax(similarities)
        score = similarities[best_idx]

        if score < 0.1:  # Threshold for "I don't know"
            return None

        # Find the specific sentence in that section
        section = self.sections[best_idx]
        sentences = re.split(r"(?<=[.!?]) +", section["text"])

        # Simple keyword matching for best sentence
        best_sent = sentences[0] if sentences else section["text"]
        max_overlap = 0
        q_tokens = set(query.lower().split())

        for sent in sentences:
            s_tokens = set(sent.lower().split())
            overlap = len(q_tokens.intersection(s_tokens))
            if overlap > max_overlap:
                max_overlap = overlap
                best_sent = sent

        return {
            "section": section["title"],
            "snippet": best_sent,
            "full_text": section["text"],
            "score": float(score),
        }

    def get_knowledge_graph(self):
        """
        Returns nodes and edges for visual graph.
        Nodes: Sections and Top Entities.
        Edges: Section -> Entity.
        """
        graph = nx.Graph()

        # Add Sections
        for sec in self.sections:
            graph.add_node(sec["title"], type="section", size=20)

            # Add top entities per section
            top_entities = [e for e in sec["entities"] if self.entity_counts[e] > 1][:5]
            for ent in top_entities:
                graph.add_node(ent, type="entity", size=10)
                graph.add_edge(sec["title"], ent)

        return graph
